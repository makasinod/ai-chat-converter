import json
import re
from docx import Document
import os


class ChatConverter:
    """Конвертує JSON чат з inputs/ у DOCX і TXT файли в outputs/"""

    def __init__(self):
        self.input_file = ""
        self.data = None
        self.chat_name = ""
        self.safe_name = ""
        self.dir_name = ""
        self.output_docx = ""
        self.output_txt = ""
        self.messages = []


    def load_json(self):
        """Зчитує JSON і бере перший елемент history[0]."""
        with open(self.input_file, "r", encoding="utf-8") as f:
            raw = json.load(f)
            self.data = raw["history"][0] if "history" in raw else raw

    def prepare_output_paths(self):
        """Створює безпечну назву, вихідну директорію та шляхи до файлів."""
        self.chat_name = self.data.get("name", os.path.splitext(os.path.basename(self.input_file))[0]).strip()
        self.safe_name = re.sub(r'[\\/*?:"<>|]', "_", self.chat_name)
        self.dir_name = os.path.join("outputs", self.safe_name)
        os.makedirs(self.dir_name, exist_ok=True)

        self.output_docx = os.path.join(self.dir_name, f"{self.safe_name}.docx")
        self.output_txt = os.path.join(self.dir_name, f"{self.safe_name}.txt")

    def parse_messages(self):
        """Витягує повідомлення з JSON."""
        data = self.data
        messages = []

        if "messages" in data:
            # OpenAI API формат
            for msg in data["messages"]:
                role = "User" if msg["role"] == "user" else "Assistant"
                text = msg.get("content", "").strip()
                if text:
                    messages.append((role, text))

        elif "mapping" in data:
            # ChatGPT Export формат
            for item in data["mapping"].values():
                message = item.get("message")
                if message and message.get("content"):
                    parts = message["content"].get("parts", [])
                    text = "\n".join(parts).strip()
                    if text:
                        author = message.get("author", {}).get("role", "unknown")
                        role = "User" if author == "user" else "Assistant"
                        messages.append((role, text))
        else:
            raise ValueError("❌ Невідомий формат JSON-файлу.")

        self.messages = messages

    def save_txt(self):
        """Зберігає чат у .txt файл."""
        text_output = [f"{role}:\n{content}\n" for role, content in self.messages]
        with open(self.output_txt, "w", encoding="utf-8") as f:
            f.write("\n".join(text_output))

    def save_docx(self):
        """Зберігає чат у .docx файл."""
        doc = Document()
        doc.add_heading(self.chat_name, level=1)

        for role, content in self.messages:
            doc.add_paragraph(f"{role}:", style="Heading 2")
            doc.add_paragraph(content)
            doc.add_paragraph("")

        doc.save(self.output_docx)

    def convert(self, input_file: str):
        """Основна логіка виконання."""
        self.input_file = input_file
        self.load_json()
        self.prepare_output_paths()
        self.parse_messages()
        self.save_txt()
        self.save_docx()

        print(f"✅ Успішно збережено:\n- {self.output_docx}\n- {self.output_txt}")
